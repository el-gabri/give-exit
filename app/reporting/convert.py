"""Markdown -> PDF / DOCX conversion.

Markdown is the canonical rendering (tested, deterministic); PDF and DOCX
are derived from it by walking a small structural subset (headings,
bullets, quotes, rules, paragraphs). One layout logic, three formats.
"""

import io
import re
from dataclasses import dataclass
from typing import Any

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
# Untrusted excerpts are backslash-escaped in the canonical Markdown so an
# uploaded file cannot inject a link, raw HTML or a code span. PDF and DOCX are
# not Markdown, so the escapes are resolved back to the literal characters here
# instead of being printed to the reader. Kept in sync with
# ``app.consumer.service._MARKDOWN_INLINE_ESCAPE_RE``.
MARKDOWN_ESCAPE_RE = re.compile(r"\\([\\`\[\]<>])")


@dataclass(frozen=True)
class Block:
    kind: str  # h1|h2|h3|bullet|quote|rule|paragraph
    text: str


def strip_markdown_escapes(text: str) -> str:
    """Resolve ``\\x`` escapes to ``x`` for the non-Markdown renderers."""
    return MARKDOWN_ESCAPE_RE.sub(r"\1", text)


def parse_blocks(markdown: str) -> list[Block]:
    blocks: list[Block] = []
    for raw in markdown.splitlines():
        line = strip_markdown_escapes(raw.rstrip())
        if not line.strip():
            continue
        if line.startswith("### "):
            blocks.append(Block("h3", line[4:]))
        elif line.startswith("## "):
            blocks.append(Block("h2", line[3:]))
        elif line.startswith("# "):
            blocks.append(Block("h1", line[2:]))
        elif line.startswith("- "):
            blocks.append(Block("bullet", line[2:]))
        elif line.startswith("> "):
            blocks.append(Block("quote", line[2:]))
        elif line.strip() == "---":
            blocks.append(Block("rule", ""))
        else:
            blocks.append(Block("paragraph", line))
    return blocks


def render_docx(markdown: str) -> bytes:
    import docx

    document = docx.Document()
    for block in parse_blocks(markdown):
        if block.kind in ("h1", "h2", "h3"):
            document.add_heading(
                BOLD_RE.sub(r"\1", block.text), level=int(block.kind[1])
            )
            continue
        if block.kind == "rule":
            continue
        style = "List Bullet" if block.kind == "bullet" else None
        paragraph = document.add_paragraph(style=style)
        if block.kind == "quote":
            paragraph.style = "Intense Quote"
        # split on **bold** spans, preserving emphasis
        cursor = 0
        for match in BOLD_RE.finditer(block.text):
            if match.start() > cursor:
                paragraph.add_run(block.text[cursor : match.start()])
            paragraph.add_run(match.group(1)).bold = True
            cursor = match.end()
        if cursor < len(block.text):
            paragraph.add_run(block.text[cursor:])

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf(markdown: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    style_map = {
        "h1": styles["Title"],
        "h2": styles["Heading2"],
        "h3": styles["Heading3"],
        "paragraph": styles["BodyText"],
        "quote": styles["Italic"],
    }

    def escape(text: str) -> str:
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return BOLD_RE.sub(r"<b>\1</b>", text)

    story: list[Any] = []
    pending_bullets: list[ListItem] = []

    def flush_bullets() -> None:
        nonlocal pending_bullets
        if pending_bullets:
            story.append(ListFlowable(pending_bullets, bulletType="bullet"))
            story.append(Spacer(1, 0.2 * cm))
            pending_bullets = []

    for block in parse_blocks(markdown):
        if block.kind == "bullet":
            pending_bullets.append(
                ListItem(Paragraph(escape(block.text), styles["BodyText"]))
            )
            continue
        flush_bullets()
        if block.kind == "rule":
            story.append(HRFlowable(width="100%"))
        else:
            story.append(Paragraph(escape(block.text), style_map[block.kind]))
            story.append(Spacer(1, 0.15 * cm))
    flush_bullets()

    buffer = io.BytesIO()
    SimpleDocTemplate(buffer, pagesize=A4, title="Notificação extrajudicial - Give Exit").build(
        story
    )
    return buffer.getvalue()
